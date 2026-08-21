"""
s12_report.py  --  Phase F deliverables

Builds the two files a reader actually opens:

  results/VALIDATION_REPORT.xlsx   every table, one sheet each
  results/VALIDATION_REPORT.docx   the written report with figures

Run s11_observed_data.py and s10_validation.py first. This script only
assembles what they produced, it computes nothing new except the small
derived quantities used in the narrative.

House style for the docx, checked by assertion at the end of this file:
no semicolons, no em dashes, no en dashes. Sentences stop on a period or
pause on a comma.
"""

from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor

ROOT = Path(__file__).resolve().parent.parent
RAW = ROOT / "data" / "raw"
TAB = ROOT / "results" / "tables"
FIG = ROOT / "results" / "figures"
REG = ROOT / "results" / "prediction_registry"
OUT = ROOT / "results"

DOX_MW = 543.52          # g/mol, doxorubicin hydrochloride free base
DOX_LIT_UM = (0.1, 10.0)


def m(name: str, metrics: pd.DataFrame) -> float:
    return float(metrics.loc[metrics["metric"] == name, "value"].iloc[0])


# ---------------------------------------------------------------------------
# spreadsheet
# ---------------------------------------------------------------------------

def build_xlsx(data: dict) -> Path:
    path = OUT / "VALIDATION_REPORT.xlsx"
    sheets = [
        ("Scorecard", data["score"]),
        ("Validation metrics", data["metrics"]),
        ("Observed reported", data["reported"]),
        ("Dose summary", data["summary"]),
        ("Per-well viability", data["viab"]),
        ("Raw absorbance", data["raw"]),
        ("Outlier audit", data["audit"]),
        ("Reconciliation", data["recon"]),
        ("Literature prior", data["prior"]),
        ("QSAR model card", data["qsar"]),
        ("Monte Carlo E1 dose range", data["e1"]),
        ("Monte Carlo E2 power", data["e2"]),
        ("Solvent confound", data["solvent"]),
    ]
    with pd.ExcelWriter(path, engine="openpyxl") as xl:
        for name, df in sheets:
            if df is None or df.empty:
                continue
            df.to_excel(xl, sheet_name=name[:31], index=False)
            ws = xl.sheets[name[:31]]
            for col in ws.columns:
                width = max(len(str(c.value)) if c.value is not None else 0
                            for c in col)
                ws.column_dimensions[col[0].column_letter].width = min(max(width + 2, 10), 52)
            ws.freeze_panes = "A2"
    return path


# ---------------------------------------------------------------------------
# document
# ---------------------------------------------------------------------------

def h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def para(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(8)
    return p


def kv_table(doc, rows, widths=(2.4, 3.6)):
    t = doc.add_table(rows=0, cols=2)
    t.style = "Light List Accent 1"
    for k, v in rows:
        cells = t.add_row().cells
        cells[0].text = str(k)
        cells[1].text = str(v)
        for r in cells[0].paragraphs[0].runs:
            r.bold = True
    for i, w in enumerate(widths):
        for row in t.rows:
            row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


def df_table(doc, df, style="Light Grid Accent 1", floatfmt="{:.2f}"):
    t = doc.add_table(rows=1, cols=len(df.columns))
    t.style = style
    for c, name in zip(t.rows[0].cells, df.columns):
        c.text = str(name)
        for r in c.paragraphs[0].runs:
            r.bold = True
    for _, row in df.iterrows():
        cells = t.add_row().cells
        for c, v in zip(cells, row):
            c.text = floatfmt.format(v) if isinstance(v, float) else str(v)
    doc.add_paragraph()
    return t


def figure(doc, path: Path, caption: str, width=6.1):
    if not path.exists():
        return
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)
    p.paragraph_format.space_after = Pt(12)


def build_docx(data: dict) -> Path:
    pred, met, score = data["pred"], data["metrics"], data["score"]
    rep = data["reported"]
    path = OUT / "VALIDATION_REPORT.docx"

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    ic50_pred = m("predicted_ic50_solvent_matched_ug_ml", met)
    ic50_obs = m("observed_ic50_lab_loglinear_ug_ml", met)
    fe = m("ic50_fold_error_solvent_matched", met)
    rmse = m("viability_rmse_pp", met)
    mae = m("viability_mae_pp", met)
    bias = m("viability_mean_bias_pp", met)
    ccc = m("lins_ccc", met)
    rho = m("spearman_rho_dose_vs_viability", met)
    v200 = m("observed_viability_at_200ppm", met)
    v400 = m("observed_viability_at_400ppm", met)
    max_inh = m("max_observed_inhibition_pct", met)
    ratio = m("lab_ic50_over_top_dose_ratio", met)
    dox = m("observed_doxorubicin_ic50_ug_ml", met)
    anova_p = m("anova_p", met)
    g5 = bool(m("gate_G5_passed", met))
    lo68, hi68 = pred["PREDICTION_1_extract_ic50_ug_per_mL"]["interval_68pct_solvent_matched"]
    dox_um = dox / DOX_MW * 1000.0

    # ---- title -----------------------------------------------------------
    t = doc.add_heading("In Silico Prediction Versus In Vitro Result", level=0)
    p = doc.add_paragraph()
    r = p.add_run("Moringa oleifera ethanolic bark extract against A549 lung "
                  "adenocarcinoma cells")
    r.italic = True
    r.font.size = Pt(12)
    para(doc, "Validation report, Phase F of the in silico study plan. "
              "Prepared 21 August 2026.", size=10)

    kv_table(doc, [
        ("Prediction sealed", f"{pred['sealed_utc']} UTC"),
        ("Sealed under commit", pred["git_commit"][:12]),
        ("SHA-256 of the sealed file",
         "7f73b66d7d33f191046f842738c3abb1395fffa482a9313f8d7be87aedfd76c2"),
        ("Assay performed by", "CCCBA Laboratory, MSU-IIT, released August 2026"),
        ("Overall verdict", "Gate G5 passed" if g5 else "Gate G5 failed"),
    ])

    # ---- summary ---------------------------------------------------------
    h(doc, "The short version", 1)
    para(doc,
         f"The computational arm predicted, before any absorbance existed, that the extract "
         f"would have an IC50 near {ic50_pred:.0f} micrograms per millilitre and that the "
         f"planned 12.5 to 200 ppm dose series would never reach 50 percent inhibition. "
         f"The laboratory measured {ic50_obs:.2f} micrograms per millilitre. That is a fold "
         f"error of {fe:.3f}, well inside the threefold band declared in advance, and it "
         f"sits inside the sealed 68 percent interval of {lo68:.0f} to {hi68:.0f}.")
    para(doc,
         f"The dose range warning also held. The laboratory added a 400 ppm point on top of "
         f"the planned series, which was more than the original protocol called for, and "
         f"viability still only fell to {v400:.2f} percent. The largest inhibition seen "
         f"anywhere in the run was {max_inh:.1f} percent. Because the response never crossed "
         f"the halfway mark, the reported IC50 was obtained by extending the regression line "
         f"past the last real data point, landing {ratio:.2f} times beyond the highest dose "
         f"actually tested.")
    para(doc,
         f"Percent viability across the five planned doses was predicted to within "
         f"{rmse:.2f} percentage points root mean square error, against a pre-declared "
         f"ceiling of 15. Dose rank order came out exactly as predicted, with a Spearman "
         f"rho of {rho:.1f}.")

    # ---- scorecard -------------------------------------------------------
    h(doc, "Scorecard", 1)
    para(doc, "Each row was written into the sealed file before the data arrived. "
              "Nothing here was chosen after seeing the result.")
    sc = score[["prediction", "criterion", "result", "verdict"]].copy()
    sc.columns = ["Sealed prediction", "Criterion", "What happened", "Verdict"]
    df_table(doc, sc)

    # ---- prediction by prediction ---------------------------------------
    h(doc, "Prediction by prediction", 1)

    h(doc, "P1, the extract IC50", 2)
    para(doc,
         f"Sealed value {ic50_pred:.1f} micrograms per millilitre, from a literature prior "
         f"built on ethanolic Moringa extracts tested against A549. Observed value "
         f"{ic50_obs:.2f}. The prediction was {fe:.3f} times the measurement, so it was "
         f"high by about 17 percent.")
    para(doc,
         "This number did not come from the machine learning model. The QSAR failed its "
         "validity gate on a scaffold split and carried roughly elevenfold error per "
         "compound on natural product frameworks, so per compound potencies were never "
         "summed into an extract IC50. The prediction rests instead on published crude "
         "extract measurements, which is a cruder method but one anchored in the same class "
         "of material, the same cell line, and the same readout.")

    h(doc, "P2, viability at each planned dose", 2)
    dose_tbl = rep[rep["sample"] == "MBEE"].sort_values("conc_ppm")
    pv = pred["PREDICTION_2_viability_at_planned_doses"]["solvent_matched_prior"]
    rows = []
    for _, r in dose_tbl.iterrows():
        key = f"{r['conc_ppm']:g}"
        key = key if key in pv else f"{float(key):.1f}"
        predicted = pv.get(key)
        rows.append({
            "Dose (ppm)": f"{r['conc_ppm']:g}",
            "Predicted %": f"{predicted:.1f}" if predicted is not None else "not predicted",
            "Observed %": f"{r['viab_mean']:.2f}",
            "SD": f"{r['viab_sd']:.2f}",
            "Residual": f"{predicted - r['viab_mean']:+.2f}" if predicted is not None else "",
        })
    df_table(doc, pd.DataFrame(rows))
    para(doc,
         f"Root mean square error {rmse:.2f} percentage points, mean absolute error "
         f"{mae:.2f}, mean signed bias {bias:+.2f}. Lin concordance correlation "
         f"coefficient {ccc:.3f}.")
    para(doc,
         "The concordance coefficient is the weakest number in this report and it deserves "
         "a plain explanation. It punishes both scatter and offset, and here the observed "
         "curve sits above 100 percent at the three lowest doses while the prediction never "
         "goes above 100 by construction. The two curves therefore agree on level and on "
         "shape at the top of the range and disagree at the bottom, which drags the "
         "coefficient down even though the absolute error stays small.")
    para(doc,
         f"Viability above 100 percent at low dose is worth noting on its own. At 12.5 ppm "
         f"the extract read {dose_tbl.iloc[0]['viab_mean']:.2f} percent, meaning the treated "
         f"wells produced more formazan than the untreated control. Two explanations fit. "
         f"Phenolic compounds can reduce the tetrazolium dye directly with no cells "
         f"involved, and low doses of plant extract can genuinely stimulate metabolism. The "
         f"cell free control that would separate these was requested before the run and was "
         f"not included, so the two cannot be told apart from this data.")

    h(doc, "P3, whether the dose range could work at all", 2)
    para(doc,
         f"The sealed claim was that the planned series would not bracket 50 percent "
         f"inhibition, with a stated falsification condition of viability at 200 ppm "
         f"dropping below 50 percent. Observed viability at 200 ppm was {v200:.2f} percent. "
         f"The claim stands.")
    para(doc,
         f"This was the finding the advisory memo was built around, and it was sent before "
         f"the assay ran. The memo put the probability of the true IC50 exceeding 200 ppm at "
         f"93.8 percent under the solvent matched prior and recommended a series running to "
         f"1600 ppm. The laboratory extended to 400 ppm, which moved in the right direction "
         f"but not far enough to capture the midpoint.")

    h(doc, "P4, dose rank order", 2)
    para(doc,
         f"Predicted a clean monotonic decrease with rho of minus one. Observed rho was "
         f"{rho:.1f} across all six doses, with no inversions. The extract behaves in an "
         f"orderly dose dependent way even though it is weak.")

    h(doc, "P5, what the statistics would show", 2)
    para(doc,
         f"The sealed claim was that no dose from 12.5 to 100 ppm would separate from the "
         f"negative control, with 200 ppm called borderline. Running the paper's own "
         f"analysis chain on the real data, the 12.5 ppm group failed the Shapiro-Wilk "
         f"normality check, so the analysis moved to the Kruskal-Wallis branch. That test "
         f"returned p equal to {anova_p:.4f}, which clears the 0.05 line, but the Dunn post "
         f"hoc comparisons separated no individual dose from the control after correction. "
         f"The sealed pattern matched on all four doses that carried a definite claim.")
    para(doc,
         "The practical reading is that the overall trend is real and no single "
         "concentration in this design can be called significantly cytotoxic on its own. "
         "That is a power problem, not a null result, and it is the reason the replicate "
         "recommendation existed.")

    h(doc, "P8, how the IC50 estimator behaves outside the tested range", 2)
    para(doc,
         f"The sealed claim was that the log linear method would place the IC50 outside the "
         f"tested range while still reporting a high coefficient of determination. The "
         f"laboratory fitted three points and reported an R squared of 0.9996 with an IC50 "
         f"of {ic50_obs:.2f}, which is {ratio:.2f} times the highest dose it actually "
         f"measured. Refitting the same log linear method on all six doses gives "
         f"{m('our_loglinear_ic50_all_doses', met):.0f} micrograms per millilitre, so the "
         f"answer moves by more than 500 units depending on which points are used.")
    para(doc,
         "A four parameter logistic fit was attempted on the same data and returned no "
         "solution, which is the correct behaviour. The curve never crosses 50 percent, so "
         "there is no midpoint to find. The straight line method returns a number anyway. "
         "That difference is the whole point of the warning.")
    para(doc,
         "The recommended wording for the manuscript is that the IC50 exceeds the highest "
         "concentration tested, with the extrapolated value reported alongside and labelled "
         "as such. Reporting 911.84 as a determination would overstate what the experiment "
         "can support.")

    # ---- figures ---------------------------------------------------------
    h(doc, "Figures", 1)
    figure(doc, FIG / "F1_predicted_vs_observed.png",
           "Figure 1. The sealed prediction against the measured dose response. "
           "Error bars are the standard deviation across the two independent trials.")
    figure(doc, FIG / "F3_ic50_regression.png",
           "Figure 2. The reported IC50 comes from extending the fitted line past the "
           "last measured dose. The shaded region is concentration that was never tested.")
    figure(doc, FIG / "F2_bland_altman.png",
           "Figure 3. Bland-Altman comparison. The prediction runs low at high dose and "
           "high at low dose, which is the signature of a slope mismatch rather than "
           "a constant offset.")

    # ---- data quality ----------------------------------------------------
    h(doc, "Data quality audit", 1)
    para(doc,
         "Raw per well absorbance was requested specifically so the reported percentages "
         "could be checked rather than taken on trust. Two things came out of that check.")

    h(doc, "The recomputation agrees", 2)
    para(doc,
         "Rebuilding percent viability from the raw optical density, using the laboratory's "
         "own formula and excluding the wells it flagged, reproduces its reported "
         "percentages to 0.59 percentage points root mean square. That confirms both the "
         "transcription and the arithmetic. The small residual is the step the report calls "
         "a correction without defining it. The laboratory's own percentages are used as "
         "primary throughout this report.")

    h(doc, "The outlier rule is aggressive and unexplained", 2)
    para(doc,
         f"The laboratory marked 27 of 78 wells as outliers, which is about a third of the "
         f"data, and attributed the decision to a Grubbs test. A Grubbs test cannot produce "
         f"that result at this design. With three wells per group the attainable range of "
         f"the test statistic barely reaches its own critical value, so running it on each "
         f"trial flags one well out of 78. Pooling all six wells per concentration flags "
         f"none. A plain rule of discarding anything more than one standard deviation from "
         f"the group mean gives 26 flags and agrees with the laboratory on 76 percent of "
         f"wells, which is close but not exact.")
    para(doc,
         "This is not a claim that the excluded wells were good data. Several of them are "
         "visibly out of line, and the trial two readings are noticeably noisier than trial "
         "one. It is a request for the actual rule to be stated, because a third of the "
         "wells being removed by an unnamed criterion is the kind of thing a reviewer will "
         "ask about, and the answer should be ready.")

    h(doc, "Solvent and dose cannot be separated in this design", 2)
    para(doc,
         "This is the most serious item in the report and it should be raised with the "
         "laboratory first. The sample preparation section describes a 10,000 microgram "
         "per millilitre stock made by dissolving 2 mg of extract in 200 microlitres of "
         "neat DMSO, with working concentrations then prepared by diluting that stock "
         "into culture medium. If the doses came straight off that stock, the final "
         "solvent fraction is fixed by the dose.")
    df_table(doc, pd.DataFrame([
        {"Dose (ug/mL)": f"{r['conc_ppm']:g}",
         "DMSO % v/v": f"{r['dmso_pct_if_direct']:.3g}",
         "Viability %": f"{r['viab_mean']:.2f}",
         "Above A549 limit": "yes" if r["over_routine_limit"] else "no"}
        for _, r in data["solvent"].iterrows()]))
    para(doc,
         "A549 is routinely held at or below 0.5 percent DMSO and the solvent alone "
         "commonly kills above about 1 percent. On these figures the top dose sits at 4 "
         "percent, eight times the routine ceiling. The three doses that showed inhibition "
         "are exactly the three that exceeded the limit, and all three doses at or below "
         "the limit read above the untreated control.")
    para(doc,
         "The correlation between solvent fraction and viability is not the point, and it "
         "is not evidence that DMSO caused the effect. Solvent fraction is an exact linear "
         "function of dose in this design, so the two variables carry identical information "
         "and cannot be told apart by any analysis of this dataset. The observed inhibition "
         "is consistent with extract activity, with solvent toxicity, or with any mixture "
         "of the two.")
    para(doc,
         "Three questions resolve it. Was an intermediate dilution in medium used before "
         "dosing, because if so the real solvent fraction may be far lower and the concern "
         "mostly goes away. What was the final percentage by volume of DMSO in the top dose "
         "well. And can the vehicle control absorbances be supplied, given that the "
         "viability formula in the report names an average vehicle absorbance term while "
         "the raw data sheet carries only negative control and blank columns.")
    figure(doc, FIG / "F4_solvent_confound.png",
           "Figure 4. Viability by dose, with the implied DMSO percentage printed inside "
           "each bar. Red bars exceed the routine A549 solvent ceiling of 0.5 percent.")

    h(doc, "The positive control ran weak", 2)
    para(doc,
         f"Doxorubicin returned an IC50 of {dox:.2f} micrograms per millilitre, which is "
         f"about {dox_um:.0f} micromolar. Published values against A549 sit between "
         f"{DOX_LIT_UM[0]} and {DOX_LIT_UM[1]} micromolar depending on exposure time and "
         f"readout, so this run came in roughly five times weaker than the top of that "
         f"range. The control still did its job of showing the assay could detect killing, "
         f"and the report is right to say so. The offset matters for a different reason.")
    para(doc,
         "The treatment exposure time is not stated anywhere in the laboratory report. The "
         "document gives the seeding density at 1.8 by ten to the fourth cells per well, an "
         "adhesion period of at least 24 hours, a 4 hour dye incubation, and a 1 hour "
         "solubilisation stand. The 4 hour figure is the formazan development step, since "
         "it is timed from the moment the dye solution was added to wells that already "
         "contained treated cells. What is missing is the gap between dosing and dye "
         "addition, which is the actual drug exposure and is normally 24, 48, or 72 hours. A short "
         "exposure would explain a weak doxorubicin result, and it would also mean the "
         "extract was given less time to act than published comparisons assume. This should "
         "be confirmed with the laboratory before the manuscript is finalised, because it "
         "changes how the potency compares to the literature.")

    h(doc, "One sentence in the laboratory analysis is not supported by the table", 2)
    para(doc,
         "The analysis section closes with the statement that the extract demonstrated "
         "inhibitory activity against A549 cells at less than 50 ug/ml. As written, that "
         "is contradicted by the report's own Table 1. Inhibitory activity means viability "
         "below the untreated control. At 50, 25 and 12.5 micrograms per millilitre the "
         "measured viability was 103.22, 107.32 and 120.30 percent, all above it. Converted "
         "to inhibition those doses give minus 3.22, minus 7.32 and minus 20.30 percent, "
         "which is stimulation rather than inhibition.")
    para(doc,
         "The unit printed is ug/ml, and that unit is used as a concentration everywhere "
         "else in the document, including for every dose in the same table. The plain "
         "reading is therefore a concentration, and on that reading the claim fails.")
    para(doc,
         "Two corrections would each make the sentence true, which is worth noting because "
         "it shows the problem is wording rather than a disagreement about the data. "
         "Changing the inequality to greater than 50 ug/ml is a single character and matches "
         "the result, since inhibition appears only at 100, 200 and 400 micrograms per "
         "millilitre. Reading the figure as a percentage instead, meaning inhibition never "
         "exceeded 50 percent, is also true, because the maximum reached anywhere was 34.96 "
         "percent, but that requires replacing the entire unit rather than one symbol.")
    para(doc,
         "That two different corrections both repair the sentence is the reason it cannot be "
         "left as an interpretation. The laboratory should state which was meant. Whichever "
         "it is, the underlying result is the same. Inhibition did not exceed 50 percent at "
         "any concentration tested, and no inhibition at all was observed below 100 "
         "micrograms per millilitre.")

    # ---- what it means ---------------------------------------------------
    h(doc, "What this means for the manuscript", 1)
    para(doc,
         f"The extract is weak against A549. At {ic50_obs:.0f} micrograms per millilitre the "
         f"IC50 is roughly nine times above the 100 ppm cutoff the National Cancer Institute "
         f"uses to call a crude extract worth pursuing, and the laboratory's own analysis "
         f"section says the same. That is a real result and it is publishable. It is not the "
         f"result the study was hoping for.")
    para(doc,
         "The honest framing is that this is a negative finding delivered with unusual "
         "rigour. The prediction was sealed and hashed before the data existed, the dose "
         "range problem was flagged in advance and then happened exactly as described, and "
         "the analysis of the real data was run through the same code as the simulation. "
         "Most undergraduate work cannot say any of that.")
    para(doc,
         "Three things should change in the write up. The IC50 should be reported as greater "
         "than 400 micrograms per millilitre with the extrapolated figure shown separately "
         "and labelled. The sentence about activity at less than 50 ug/ml should be "
         "rewritten to say plainly that inhibition did not exceed 50 percent at any "
         "concentration tested. The absence of a vehicle control and a cell free interference control "
         "should be stated as limitations, since without them the above 100 percent "
         "readings at low dose cannot be explained.")

    h(doc, "What the model got wrong", 1)
    para(doc,
         f"The prediction was too flat at the bottom of the curve. It expected 99.6 percent "
         f"viability at 12.5 ppm and the assay read 120.30, a miss of {abs(-20.70):.1f} "
         f"points, which is where most of the total error comes from. The generative model "
         f"caps viability at 100 percent, so it structurally cannot produce the low dose "
         f"stimulation the data shows. That is a fixable modelling choice rather than a "
         f"failure of the potency estimate.")
    para(doc,
         "The pooled prior, which was reported as the optimistic bound, predicted 304 "
         "micrograms per millilitre and was wrong by a factor of three. The solvent matched "
         "prior was the one nominated in advance as decision relevant, and it was the one "
         "that held. Choosing it before seeing the data is the part that mattered.")

    h(doc, "Limitations", 1)
    for text in [
        "No published A549 data exists for Moringa bark, so the prior was built from leaf "
        "extracts. The agreement found here is therefore partly fortunate and should not be "
        "read as evidence that leaf and bark are interchangeable.",
        "The comparison rests on two independent trials. That is thin, and the standard "
        "deviations across trials run as high as 21 percentage points.",
        "No quantitative composition data was supplied for the batch tested, so the "
        "docking and network results describe Moringa chemistry as reported in the "
        "literature rather than what was in this material.",
        "The treatment exposure time is unknown, which limits how far the potency can be "
        "compared against published values.",
        "The laboratory applied an undocumented correction and an unstated outlier rule "
        "before reporting percentages.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(text).font.size = Pt(10.5)

    h(doc, "Files", 1)
    kv_table(doc, [
        ("Spreadsheet", "results/VALIDATION_REPORT.xlsx"),
        ("Sealed prediction", "results/prediction_registry/prediction_v1.json"),
        ("Hash ledger", "results/prediction_registry/REGISTRY.md"),
        ("Dose range advisory", "results/prediction_registry/DOSE_RANGE_MEMO.md"),
        ("Full statistics output", "results/tables/F3_statistics.txt"),
        ("Ingest code", "src/s11_observed_data.py"),
        ("Validation code", "src/s10_validation.py"),
        ("Methods", "docs/METHODS.docx"),
    ], widths=(2.0, 4.2))

    doc.save(path)
    return path


# ---------------------------------------------------------------------------

def check_style(path: Path) -> None:
    """House style gate. Fails loudly rather than shipping a bad file."""
    from docx import Document as D
    d = D(str(path))
    text = "\n".join(p.text for p in d.paragraphs)
    for tbl in d.tables:
        for row in tbl.rows:
            for c in row.cells:
                text += "\n" + c.text
    bad = {";": text.count(";"), "em dash": text.count("—"),
           "en dash": text.count("–")}
    if any(bad.values()):
        raise AssertionError(f"house style violated: {bad}")
    print(f"  style check passed, no semicolons or long dashes")


def main() -> None:
    def load(name):
        p = TAB / name
        return pd.read_csv(p) if p.exists() else pd.DataFrame()

    data = {
        "pred": json.loads((REG / "prediction_v1.json").read_text(encoding="utf-8")),
        "metrics": load("F2_validation_metrics.csv"),
        "score": load("F2_prediction_scorecard.csv"),
        "summary": load("F1_dose_summary.csv"),
        "viab": load("F1_observed_viability.csv"),
        "audit": load("F1_grubbs_audit.csv"),
        "recon": load("F1_reconciliation.csv"),
        "prior": load("literature_prior.csv"),
        "qsar": load("qsar_model_card.csv"),
        "e1": load("E1_dose_range_adequacy.csv"),
        "e2": load("E2_power.csv"),
        "reported": pd.read_csv(RAW / "mtt_observed_reported.csv"),
        "solvent": load("F4_solvent_confound.csv"),
        "raw": pd.read_csv(RAW / "mtt_observed_absorbance.csv"),
    }

    print("=" * 78)
    print("PHASE F -- building deliverables")
    print("=" * 78)
    x = build_xlsx(data)
    print(f"\n  spreadsheet -> {x.relative_to(ROOT)}")
    d = build_docx(data)
    print(f"  document    -> {d.relative_to(ROOT)}")
    check_style(d)


if __name__ == "__main__":
    main()
